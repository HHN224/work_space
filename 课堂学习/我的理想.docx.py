from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('我的理想')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 正文内容
content = (
    '每个人都有自己的梦想，而我的理想是成为一名优秀的软件工程师，为世界计算机事业的发展贡献自己的力量。\n\n'
    '从小我就对计算机充满了浓厚的兴趣。每当看到那些神奇的程序在屏幕上运行，帮助人们解决各种复杂的问题，'
    '我的心中便充满了敬佩与向往。我渴望有一天，自己也能写出改变世界的代码，让科技更好地服务人类。\n\n'
    '软件工程师不仅仅是写代码的人，更是用技术连接世界、创造未来的桥梁。我希望通过自己的努力，'
    '开发出更加智能、高效的软件系统，推动人工智能、云计算等前沿技术的发展，让偏远地区的孩子也能通过互联网接触知识，'
    '让医疗资源通过智能系统惠及更多患者。\n\n'
    '为了实现这个理想，我现在就努力学习数学和英语，打下坚实的基础。我相信，只要坚持不懈地奋斗，'
    '终有一天，我能在计算机科学的广阔天地里，留下属于自己的印记，为世界的发展添砖加瓦。'
)

paragraphs = content.split('\n\n')
for text in paragraphs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    p.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 保存文档
doc.save('/Users/hhn224/Desktop/my_project/work_space/课堂学习/我的理想.docx')
print('文档已生成：我的理想.docx')
